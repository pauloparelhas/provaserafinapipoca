from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

def heading(doc, text, level=1, rgb=None):
    p = doc.add_heading(text, level=level)
    if rgb:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*rgb)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

AZUL = (0x2E, 0x86, 0xAB)
ROXO = (0x56, 0x4D, 0x65)
VERDE = (0x27, 0xAE, 0x60)
AMARELO = (0xF3, 0x9C, 0x12)
VERMELHO = (0xE8, 0x3F, 0x3F)
CINZA = (0x95, 0xA5, 0xA6)
ESCURO = (0x2C, 0x3E, 0x50)

# ---- CAPA ----
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('PROJETO SERAFINA')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(*AZUL)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run('Plataforma Gamificada de Preparacao para Provas')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(*ROXO)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Maple Bear  |  Y2 Elementary  |  1o Trimestre 2026\n')
m.add_run('Crianca: 7 anos\n')
m.add_run('Documento criado em: ' + datetime.date.today().strftime('%d/%m/%Y'))

doc.add_page_break()

# ---- 1. VISAO GERAL ----
heading(doc, '1. VISAO GERAL DO PROJETO', 1, AZUL)
doc.add_paragraph(
    'O Projeto Serafina e uma iniciativa de preparacao para as Avaliacoes Formais (AFs) '
    'do 1o trimestre da aluna do Y2 Elementary da Maple Bear. Consiste na criacao de '
    'ferramentas digitais gamificadas e interativas, organizadas por assunto, '
    'para que a crianca aprenda brincando com protagonismo, desafios progressivos e feedback imediato.'
)

heading(doc, '1.1 Filosofia do Projeto', 2, ROXO)
for item in [
    'A crianca e protagonista do aprendizado',
    'Gamificado: pontuacao, niveis, conquistas e estrelas',
    'Feedback imediato: acerto (verde, estrelas) e erro (vermelho, dica + tente novamente)',
    'Progressao de dificuldade: do simples ao complexo (niveis)',
    'Desafiador, mas nunca frustrante',
    'Visual atrativo, colorido, personagem Serafina como mascote',
    'Ferramentas acessiveis por link (HTML puro, sem instalacao)',
]:
    bullet(doc, item)

doc.add_page_break()

# ---- 2. CALENDARIO ----
heading(doc, '2. AVALIACOES FORMAIS - CALENDARIO', 1, AZUL)
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
for i, txt in enumerate(['Data', 'Disciplina', 'Status']):
    c = t2.rows[0].cells[i]
    c.text = txt
    for p in c.paragraphs:
        for r in p.runs: r.bold = True

for row_data in [
    ('10/03/2026', 'AF Lingua Portuguesa', 'PRIORITARIA'),
    ('18/03/2026', 'AF E.L.A. (Ingles)', 'A PLANEJAR'),
    ('26/03/2026', 'AF Historia', 'A PLANEJAR'),
]:
    row = t2.add_row()
    for i, v in enumerate(row_data):
        row.cells[i].text = v

doc.add_paragraph()
doc.add_page_break()

# ---- 3. CONTEUDOS ----
heading(doc, '3. CONTEUDOS POR DISCIPLINA', 1, AZUL)

heading(doc, '3.1 Lingua Portuguesa (AF: 10/03)', 2, VERMELHO)
lp = [
    ('LP-01', 'Genero textual - Diario', 'O que e um diario, suas caracteristicas, como e escrito'),
    ('LP-02', 'Livro: O Diario Escondido de Serafina', 'Personagem, enredo, cenarios, aventuras'),
    ('LP-03', 'Perfil da personagem Serafina', 'Quem e, onde vive, passatempos, esconderijo atual'),
    ('LP-04', 'Alfabeto maiusculo e minusculo', 'Reconhecer, parear e escrever letras nas duas formas'),
    ('LP-05', 'Ordem alfabetica', 'Ordenar palavras pela 1a e 2a letra (brincadeiras)'),
    ('LP-06', 'Separacao silabica', 'Dividir palavras em silabas corretamente [JA EXISTE]'),
    ('LP-07', 'Uso da letra maiuscula', 'Regras: inicio de frase, nomes proprios, etc.'),
    ('LP-08', 'Escrita independente: Meu Esconderijo', 'Producao textual: titulo, local, objetos, atividades'),
]
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
for i, txt in enumerate(['Codigo', 'Conteudo', 'Descricao']):
    c = t3.rows[0].cells[i]
    c.text = txt
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
for row_data in lp:
    row = t3.add_row()
    for i, v in enumerate(row_data): row.cells[i].text = v

doc.add_paragraph()
heading(doc, '3.2 E.L.A. - English Language Arts (AF: 18/03)', 2, VERDE)
ela = [
    ('ELA-01', 'Acrostic Poem', 'Criar e identificar poemas acrosticos'),
    ('ELA-02', 'Bucket Filler vs Bucket Dipper', 'Acoes que enchem ou esvaziam o balde do outro'),
    ('ELA-03', 'Community Helpers', 'Profissoes que ajudam a comunidade'),
    ('ELA-04', 'Sense of Community', 'Acoes: help, respect, empathy'),
    ('ELA-05', 'Family Vocabulary', 'Vocabulario de familia em ingles'),
    ('ELA-06', 'Reading & Interpreting Texts', 'Compreensao leitora em ingles'),
]
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
for i, txt in enumerate(['Codigo', 'Conteudo', 'Descricao']):
    c = t4.rows[0].cells[i]
    c.text = txt
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
for row_data in ela:
    row = t4.add_row()
    for i, v in enumerate(row_data): row.cells[i].text = v

doc.add_paragraph()
heading(doc, '3.3 Historia (AF: 26/03)', 2, AMARELO)
doc.add_paragraph('Unidade 1: Quanto o tempo o tempo tem? - Conteudo a detalhar apos materiais de sala.')

doc.add_page_break()

# ---- 4. FERRAMENTAS ----
heading(doc, '4. FERRAMENTAS DIGITAIS - PLANEJAMENTO', 1, AZUL)

ferramentas = [
    {
        'cod': 'LP-06', 'nome': 'Separacao Silabica', 'status': 'CONCLUIDA',
        'tipo': 'Dashboard interativo',
        'mecanica': 'Exercicios de separacao silabica com feedback visual e animacoes',
        'link': 'Ja disponivel - ver arquivo links de ferramentas.docx',
    },
    {
        'cod': 'LP-05', 'nome': 'Ordem Alfabetica', 'status': 'PLANEJADA',
        'tipo': 'Jogo drag and drop com niveis',
        'mecanica': 'Nivel 1: 5 letras | Nivel 2: 10 letras | Nivel 3: palavras de brincadeiras (1a letra) | Nivel 4: palavras com mesma 1a letra (2a letra). Feedback: verde/vermelho, botao de dica, tela de parabens.',
        'link': 'A gerar',
    },
    {
        'cod': 'LP-04', 'nome': 'Alfabeto Maiusculo e Minusculo', 'status': 'PLANEJADA',
        'tipo': 'Teoria em cards + Drag and Drop',
        'mecanica': 'Cards de teoria rapida → arrastar maiusculas para um lado e minusculas para outro → fase bonus: parear letra maiuscula com sua minuscula',
        'link': 'A gerar',
    },
    {
        'cod': 'LP-07', 'nome': 'Uso da Letra Maiuscula', 'status': 'PLANEJADA',
        'tipo': 'Teoria visual + Quiz interativo',
        'mecanica': 'Regras em cards animados → identificar erros em frases → corrigir textos clicando na letra certa',
        'link': 'A gerar',
    },
    {
        'cod': 'LP-02/LP-03', 'nome': 'Personagem Serafina', 'status': 'AGUARDANDO LIVRO',
        'tipo': 'Perfil interativo + Quiz',
        'mecanica': 'Perfil visual da Serafina → quiz de interpretacao textual → completar o perfil da personagem',
        'link': 'A gerar apos disponibilizacao do livro',
    },
    {
        'cod': 'LP-01', 'nome': 'Genero Textual: Diario', 'status': 'PLANEJADA',
        'tipo': 'Teoria visual + Identificacao de elementos',
        'mecanica': 'O que e um diario → caracteristicas em cards → identificar elementos em um diario de exemplo → mini-diario interativo',
        'link': 'A gerar',
    },
    {
        'cod': 'LP-08', 'nome': 'Meu Esconderijo (Escrita)', 'status': 'PLANEJADA',
        'tipo': 'Guia de producao textual',
        'mecanica': 'Template estruturado passo a passo: titulo → descricao do local → objetos que levaria → atividades que faz. Dicas visuais em cada etapa.',
        'link': 'A gerar',
    },
]

for f in ferramentas:
    heading(doc, f'[{f["cod"]}] {f["nome"]}', 3, ESCURO)
    tf = doc.add_table(rows=4, cols=2)
    tf.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ('Status', f['status']),
        ('Tipo', f['tipo']),
        ('Mecanica de Jogo', f['mecanica']),
        ('Link', f['link']),
    ]):
        tf.rows[i].cells[0].text = label
        tf.rows[i].cells[1].text = value
        for p in tf.rows[i].cells[0].paragraphs:
            for r in p.runs: r.bold = True
    doc.add_paragraph()

doc.add_page_break()

# ---- 5. EQUIPE ----
heading(doc, '5. EQUIPE MULTIDISCIPLINAR DE AGENTES', 1, AZUL)
agentes = [
    ('Coordenador Pedagogico',
     'Responsavel geral: define escopo, aprova conteudo, garante alinhamento com a prova e adequacao para 7 anos'),
    ('Especialista em Ferramentas Educacionais',
     'Define mecanicas de gamificacao, UX infantil, referencias de edtech, padroes de interacao'),
    ('Gramatico/Pedagogo de LP',
     'Valida conteudo linguistico: silabas, ordem alfabetica, maiusculas/minusculas, regras para Y2'),
    ('Desenvolvedor de Jogos Educativos',
     'Cria as ferramentas em HTML/CSS/JavaScript: animacoes, pontuacao, niveis, feedback visual'),
    ('Revisor de Desenvolvimento',
     'Testa ferramentas: bugs, fluxo de jogo, dificuldade, mensagens de feedback, cross-browser'),
    ('Gerenciador de Projeto (IA)',
     'Mantem MEMORY.md atualizado, documenta etapas, controla backlog, salva progresso a cada 5min'),
]
for nome, desc in agentes:
    bullet(doc, desc, bold_prefix=nome + ': ')

doc.add_page_break()

# ---- 6. ETAPAS ----
heading(doc, '6. ETAPAS DO PROJETO', 1, AZUL)

etapas = [
    ('ETAPA 0 - FUNDACAO', 'CONCLUIDA', VERDE, [
        'Leitura e analise dos materiais da escola (8 imagens)',
        'Identificacao de todas as disciplinas e conteudos',
        'Criacao deste documento de planejamento',
        'Criacao da estrutura de memoria persistente (MEMORY.md)',
        'Mapeamento das ferramentas existentes (separacao silabica)',
    ]),
    ('ETAPA 1 - LP: Ordem Alfabetica', 'AGUARDANDO APROVACAO', AMARELO, [
        'Design do jogo: niveis, feedback, dicas, tela de parabens',
        'Banco de dados: letras e palavras de brincadeiras',
        'Desenvolvimento HTML/CSS/JS',
        'Teste e validacao pedagogica',
        'Publicacao e registro do link',
    ]),
    ('ETAPA 2 - LP: Alfabeto Maiusculo e Minusculo', 'AGUARDANDO APROVACAO', AMARELO, [
        'Design: teoria em cards + drag and drop',
        'Banco de dados: pares de letras',
        'Desenvolvimento HTML/CSS/JS',
        'Teste e publicacao',
    ]),
    ('ETAPA 3 - LP: Uso da Letra Maiuscula', 'AGUARDANDO APROVACAO', AMARELO, [
        'Levantamento das regras adequadas ao Y2',
        'Design: teoria + identificacao em frases',
        'Desenvolvimento, teste e publicacao',
    ]),
    ('ETAPA 4 - LP: Personagem Serafina', 'AGUARDANDO LIVRO', CINZA, [
        'Leitura do livro O Diario Escondido de Serafina',
        'Extracao de informacoes do perfil da personagem',
        'Design: perfil interativo + quiz de interpretacao',
        'Desenvolvimento e publicacao',
    ]),
    ('ETAPA 5 - LP: Genero Textual Diario', 'AGUARDANDO APROVACAO', AMARELO, [
        'Cards de teoria: o que e um diario e suas caracteristicas',
        'Exercicios de identificacao de elementos',
        'Mini-diario interativo',
    ]),
    ('ETAPA 6 - LP: Meu Esconderijo', 'AGUARDANDO APROVACAO', AMARELO, [
        'Template estruturado de producao textual',
        'Guia passo a passo com dicas visuais',
    ]),
    ('ETAPA 7 - ELA (Ingles)', 'A PLANEJAR', CINZA, [
        'Levantar materiais (Collection Folder, workbook)',
        'Planejar 6 ferramentas (ELA-01 a ELA-06)',
        'Desenvolvimento sequencial por conteudo',
    ]),
    ('ETAPA 8 - Historia', 'A PLANEJAR', CINZA, [
        'Levantar materiais da unidade Quanto o tempo o tempo tem?',
        'Planejar e desenvolver ferramentas',
    ]),
]

for nome, status, cor, subitems in etapas:
    heading(doc, nome, 2, (0x27, 0x60, 0x87))
    sp = doc.add_paragraph()
    rr = sp.add_run('Status: ' + status)
    rr.bold = True
    rr.font.color.rgb = RGBColor(*cor)
    for item in subitems:
        bullet(doc, item)

doc.add_page_break()

# ---- 7. REGISTRO DE LINKS ----
heading(doc, '7. REGISTRO DE FERRAMENTAS GERADAS', 1, AZUL)
doc.add_paragraph('Esta secao sera atualizada a cada nova ferramenta concluida.')
tl = doc.add_table(rows=1, cols=4)
tl.style = 'Table Grid'
for i, txt in enumerate(['Codigo', 'Ferramenta', 'Link', 'Data']):
    c = tl.rows[0].cells[i]
    c.text = txt
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
row = tl.add_row()
row.cells[0].text = 'LP-06'
row.cells[1].text = 'Separacao Silabica'
row.cells[2].text = 'Ver arquivo links de ferramentas.docx'
row.cells[3].text = 'Anterior a 07/03/2026'

doc.add_page_break()

# ---- 8. NOTAS ----
heading(doc, '8. NOTAS, DECISOES E HISTORICO', 1, AZUL)
notas = [
    '07/03/2026 | Projeto iniciado. Materiais analisados: 8 imagens de instrucoes da escola.',
    '07/03/2026 | Livro O Diario Escondido de Serafina ainda nao disponivel na pasta.',
    '07/03/2026 | Ferramenta separacao silabica ja existe (ver links de ferramentas.docx).',
    '07/03/2026 | Prioridade: AF Lingua Portuguesa (10/03) - 7 ferramentas a criar.',
    '07/03/2026 | Tecnologia: HTML/CSS/JavaScript puro - funciona como link sem servidor.',
    '07/03/2026 | Estilo: gamificado, Serafina como mascote, cores vibrantes, drag and drop.',
]
for nota in notas:
    bullet(doc, nota)

out = r'C:\Users\paulo\OneDrive\td junto outlook hotmail\MAPLE BEAR\provas\planejamento\PROJETO_SERAFINA_Planejamento.docx'
doc.save(out)
print('OK: ' + out)
