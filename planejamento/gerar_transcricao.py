from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

def h(doc, text, level=1, rgb=None):
    p = doc.add_heading(text, level=level)
    if rgb:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*rgb)
    return p

def para(doc, text, bold=False, rgb=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if rgb: r.font.color.rgb = RGBColor(*rgb)
    if size: r.font.size = Pt(size)
    return p

def bul(doc, text, bold_pre=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_pre:
        rb = p.add_run(bold_pre); rb.bold = True
    p.add_run(text)
    return p

AZUL = (0x2E, 0x86, 0xAB)
ROXO = (0x56, 0x4D, 0x65)
VERDE = (0x27, 0xAE, 0x60)
AMARELO = (0xD3, 0x54, 0x00)
VERMELHO = (0xC0, 0x39, 0x2B)
CINZA = (0x60, 0x60, 0x60)

# ---- CAPA ----
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('TRANSCRICAO DOS MATERIAIS ESCOLARES')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(*AZUL)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run('Projeto Serafina  |  Maple Bear Y2 Elementary  |  1o Trimestre 2026')
r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(*ROXO)

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Transcricao fiel das imagens enviadas pela escola\n')
m.add_run('Gerado em: ' + datetime.date.today().strftime('%d/%m/%Y'))

doc.add_page_break()

# ====================================================
# SECAO 1 - INFORMATIVO GERAL
# ====================================================
h(doc, '1. INFORMATIVO GERAL DAS AFs - 1o TRIMESTRE 2026', 1, AZUL)
para(doc, 'Fonte: image8.png | Arquivo: 00_COORDENACAO/informativo_AFs_1trimestre_2026.png',
     rgb=CINZA, size=9)
doc.add_paragraph()

h(doc, 'AF LINGUA PORTUGUESA - 10/03', 2, VERMELHO)
para(doc, 'Unidade 1 - Sabe quem sou eu?', bold=True)

h(doc, 'Conteudos:', 3, CINZA)
for item in [
    'Genero textual - Diario',
    'Livro literario: O Diario Escondido de Serafina',
    'O alfabeto maiusculo e minusculo',
    'Ordem alfabetica',
    'Separacao silabica',
    'Uso da letra maiuscula',
]:
    bul(doc, item)

h(doc, 'Materiais para estudo:', 3, CINZA)
for item in [
    'Caderno',
    'Pasta Collection Folder',
    'Recursos compartilhados no drive da turma (folder) - OPCIONAL',
]:
    bul(doc, item)

doc.add_paragraph()
h(doc, 'AF E.L.A. (ENGLISH) - 18/03', 2, VERDE)
para(doc, 'Unit 1 - Building A Classroom Community', bold=True)
for item in [
    'Acrostic poem',
    'The difference between bucket filler and bucket dipper',
    'Community helpers',
    'Actions that show sense of community (help, respect, empathy etc.)',
]:
    bul(doc, item)

para(doc, 'Unit 2 - Family and Friends Customs', bold=True)
for item in [
    'Family vocabulary',
    'Reading and interpreting texts',
]:
    bul(doc, item)

h(doc, 'Materials:', 3, CINZA)
for item in [
    'Collection Folder',
    'E.L.A. workbook',
    'Recursos compartilhados no drive da turma (folder)',
]:
    bul(doc, item)

doc.add_paragraph()
h(doc, 'AF HISTORIA - 26/03', 2, (0xF3, 0x9C, 0x12))
para(doc, 'Unidade 1 - Quanto o tempo o tempo tem?', bold=True)
h(doc, 'Materiais:', 3, CINZA)
for item in [
    'Recursos compartilhados no drive da turma (folder) - OPCIONAL',
    'Collections folders e caderno',
]:
    bul(doc, item)

doc.add_page_break()

# ====================================================
# SECAO 2 - LP: ORDEM ALFABETICA
# ====================================================
h(doc, '2. LINGUA PORTUGUESA - ORDEM ALFABETICA', 1, AZUL)
para(doc, 'Fontes: image4.png e image5.png | Pasta: 01_LINGUA_PORTUGUESA/materiais_escola/',
     rgb=CINZA, size=9)
doc.add_paragraph()

h(doc, '2.1 Instrucoes da Atividade (image4.png)', 2, ROXO)
para(doc, 'Aspectos linguisticos - Ordem alfabetica', bold=True)
para(doc, 'Objetivo:', bold=True)
para(doc, 'Desenvolver a compreensao da ordem alfabetica e ampliar o repertorio de palavras.')

para(doc, 'Como realizar a atividade:', bold=True)
for item in [
    'Primeiramente, o(a) estudante devera recortar as tarjetas com os nomes das brincadeiras.',
    'Em seguida, devera observar atentamente a letra inicial de cada palavra.',
    'Organizara os nomes em ordem alfabetica, considerando a sequencia correta das letras.',
    'Caso duas palavras iniciem com a mesma letra, devera observar a segunda letra para definir a ordem.',
    'Apos organizar, o(a) estudante devera registrar a ordem no caderno.',
]:
    bul(doc, item)

doc.add_paragraph()
h(doc, '2.2 Exercicio com Nomes de Brincadeiras (image5.png)', 2, ROXO)
para(doc, 'Atividade 1: Organizar em ordem alfabetica', bold=True)
para(doc, 'Palavras fornecidas (tarjetas para recortar):')

palavras = [
    'pular corda', 'bambole', 'queimada',
    'danca das cadeiras', 'pega-pega', 'policia e ladrao',
    'morto-vivo', 'estatua', 'mimica',
    'corrida de batata', 'bolhas de sabao', 'cabra-cega',
    'alerta', 'passa anel', 'amarelinha',
]
for p in palavras:
    bul(doc, p)

doc.add_paragraph()
para(doc, 'Atividade 2: Completar o alfabeto', bold=True)
para(doc, 'Instrucao: Preencha o alfabeto com as letras que faltam para confirmar a organizacao que voce realizou.')
para(doc, 'O alfabeto apresentado tinha as seguintes letras faltando (a serem preenchidas pelo aluno):')
para(doc, 'F, H, I, K, M, O, R, W, Y, Z  (e possivelmente outras - confirmar com caderno)')
doc.add_paragraph()
para(doc, 'Nota pedagogica:', bold=True, rgb=VERMELHO)
para(doc, 'Esta atividade requer que o aluno domine tanto a sequencia das letras do alfabeto quanto a ordenacao de palavras pela letra inicial - e quando ha empate, pela segunda letra.')

doc.add_page_break()

# ====================================================
# SECAO 3 - PERFIL SERAFINA
# ====================================================
h(doc, '3. LINGUA PORTUGUESA - PERFIL DA PERSONAGEM SERAFINA', 1, AZUL)
para(doc, 'Fontes: image3.png e image6.png | Pasta: 01_LINGUA_PORTUGUESA/materiais_escola/',
     rgb=CINZA, size=9)
doc.add_paragraph()

h(doc, '3.1 Instrucoes (image3.png)', 2, ROXO)
para(doc, 'Atividade 2 - Perfil da personagem Serafina', bold=True)
para(doc, 'Objetivo:', bold=True)
para(doc, 'Estimular a interpretacao textual e a identificacao de caracteristicas da personagem.')

para(doc, 'Como realizar a atividade:', bold=True)
for item in [
    'O(a) estudante devera relembrar as informacoes apresentadas na leitura sobre a personagem Serafina.',
    'Em cada quadro, respondera as perguntas com frases completas.',
    'E importante que as respostas sejam baseadas no texto trabalhado em sala.',
]:
    bul(doc, item)

para(doc, 'Esta atividade desenvolve a compreensao leitora, a organizacao de ideias e a escrita com sentido.')

doc.add_paragraph()
h(doc, '3.2 Ficha do Perfil (image6.png)', 2, ROXO)
para(doc, 'Pergunta central: O que voce ja sabe sobre a personagem Serafina?', bold=True)
doc.add_paragraph()

campos_perfil = [
    ('Em que local(is) a Serafina vive?', 'A ser preenchido com base no livro'),
    ('O que Serafina gosta de fazer (seus passatempos)?', 'A ser preenchido com base no livro'),
    ('Nomes de/da Serafina', 'A ser preenchido com base no livro'),
    ('Por que Serafina consegue ou nao se encaixa nos grupos?', 'A ser preenchido com base no livro'),
    ('O que Serafina faz em seu atual esconderijo?', 'A ser preenchido com base no livro'),
]

tf = doc.add_table(rows=1, cols=2)
tf.style = 'Table Grid'
tf.rows[0].cells[0].text = 'Campo do Perfil'
tf.rows[0].cells[1].text = 'Resposta Esperada'
for cell in tf.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True

for campo, resp in campos_perfil:
    row = tf.add_row()
    row.cells[0].text = campo
    row.cells[1].text = resp

doc.add_paragraph()
para(doc, 'IMPORTANTE: As respostas corretas dependem da leitura do livro O Diario Escondido de Serafina.',
     bold=True, rgb=VERMELHO)
para(doc, 'Assim que o livro for disponibilizado na pasta, este documento sera atualizado com as respostas.')

doc.add_page_break()

# ====================================================
# SECAO 4 - ESCRITA INDEPENDENTE
# ====================================================
h(doc, '4. LINGUA PORTUGUESA - ESCRITA INDEPENDENTE: MEU ESCONDERIJO', 1, AZUL)
para(doc, 'Fontes: image7.png, image1.png, image2.png | Pasta: 01_LINGUA_PORTUGUESA/materiais_escola/',
     rgb=CINZA, size=9)
doc.add_paragraph()

h(doc, '4.1 Instrucoes (image1.png e image2.png)', 2, ROXO)
para(doc, 'Atividade 3 - Escrita Independente: Meu Esconderijo', bold=True)
para(doc, 'Objetivo:', bold=True)
para(doc, 'Estimular a producao textual autoral, a criatividade e a organizacao do texto.')

para(doc, 'Como realizar a atividade:', bold=True)
for item in [
    'O(a) estudante devera criar um titulo para seu texto.',
    'Em seguida, descrever o local do seu esconderijo (onde fica, como e, o que tem la).',
    'Depois, escrever o que levaria para esse esconderijo.',
    'Por fim, explicar o que faz nesse espaco.',
]:
    bul(doc, item)

para(doc, 'Orientar para que:', bold=True)
for item in [
    'Utilizem frases completas.',
    'Organizem as ideias em sequencia.',
    'Releiam o texto apos escrever para verificar se faz sentido.',
]:
    bul(doc, item)

para(doc, 'Esta atividade fortalece a autonomia na escrita, amplia o vocabulario e desenvolve a criatividade.')

doc.add_paragraph()
h(doc, '4.2 Template da Atividade (image7.png)', 2, ROXO)
para(doc, 'Estrutura da folha de exercicio:', bold=True)

template_campos = [
    'Titulo: _______________________________________________',
    'Descricao do local do esconderijo: _____________________',
    'O que voce levaria para seu esconderijo? ________________',
    'O que voce faz no seu esconderijo? ____________________',
]
for campo in template_campos:
    bul(doc, campo)

doc.add_paragraph()
para(doc, 'Nota pedagogica:', bold=True, rgb=VERMELHO)
para(doc, 'A ferramenta digital para esta atividade devera guiar a crianca pelos 4 campos de forma '
         'sequencial e visual, com dicas em cada etapa e uma previa do texto sendo formado.')

doc.add_page_break()

# ====================================================
# SECAO 5 - NOTA SOBRE CONTEUDOS PENDENTES
# ====================================================
h(doc, '5. CONTEUDOS PENDENTES DE MATERIALIZACAO', 1, AZUL)

h(doc, '5.1 Conteudos sem material visual ainda', 2, ROXO)
pendentes = [
    ('LP-01 - Genero Textual Diario',
     'Nao ha imagem especifica. Conteudo a ser detalhado com base no caderno e orientacoes de sala.'),
    ('LP-04 - Alfabeto Maiusculo e Minusculo',
     'Instrucao mencionada na image3.png mas sem exercicio visual separado. Ver caderno.'),
    ('LP-06 - Separacao Silabica',
     'Ja possui ferramenta digital. Ver links de ferramentas.docx.'),
    ('LP-07 - Uso da Letra Maiuscula',
     'Conteudo mencionado no informativo (image8.png) mas sem exercicio visual. Ver caderno.'),
]

for nome, obs in pendentes:
    bul(doc, obs, bold_pre=nome + ': ')

doc.add_paragraph()
h(doc, '5.2 Livro Literario - O Diario Escondido de Serafina', 2, ROXO)
para(doc,
    'O livro ainda nao foi disponibilizado na pasta do projeto. '
    'Ele e essencial para as atividades LP-02 e LP-03 (perfil e interpretacao da Serafina). '
    'Assim que disponibilizado, sera feita a leitura e extracao de conteudo para:'
)
for item in [
    'Respostas do perfil da personagem',
    'Perguntas de interpretacao textual',
    'Elementos do diario como genero textual',
    'Vocabulario relevante para as demais atividades',
]:
    bul(doc, item)

doc.add_page_break()

# ====================================================
# SECAO 6 - INDICE DE ARQUIVOS
# ====================================================
h(doc, '6. INDICE DE ARQUIVOS DO PROJETO', 1, AZUL)
para(doc, 'Estrutura de pastas criada em 07/03/2026:', bold=True)

estrutura = [
    ('00_COORDENACAO/', 'Planejamento geral, taxonomia, informativo das AFs'),
    ('  informativo_AFs_1trimestre_2026.png', 'Print do informativo com calendario e conteudos'),
    ('  PROJETO_SERAFINA_Planejamento.docx', 'Documento master de planejamento do projeto'),
    ('  preparacao_para_prova_original.docx', 'Arquivo original da escola'),
    ('  taxonomia/', 'Mapas de organizacao e taxonomia do projeto'),
    ('01_LINGUA_PORTUGUESA/', 'Tudo sobre a AF de Lingua Portuguesa (10/03)'),
    ('  materiais_escola/', 'Imagens e materiais enviados pela escola'),
    ('  ferramentas/', 'Ferramentas digitais geradas para LP'),
    ('02_ELA_INGLES/', 'Tudo sobre a AF de Ingles/ELA (18/03)'),
    ('  materiais_escola/', 'Materiais da escola para ELA'),
    ('  ferramentas/', 'Ferramentas digitais geradas para ELA'),
    ('03_HISTORIA/', 'Tudo sobre a AF de Historia (26/03)'),
    ('04_LIVRO_SERAFINA/', 'Livro O Diario Escondido de Serafina (aguardando)'),
    ('05_LINKS_FERRAMENTAS/', 'Arquivo com todos os links das ferramentas geradas'),
    ('planejamento/', 'Scripts e versoes do documento de planejamento'),
]

tind = doc.add_table(rows=1, cols=2)
tind.style = 'Table Grid'
tind.rows[0].cells[0].text = 'Caminho'
tind.rows[0].cells[1].text = 'Descricao'
for cell in tind.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True
for caminho, desc in estrutura:
    row = tind.add_row()
    row.cells[0].text = caminho
    row.cells[1].text = desc

out = r'C:\Users\paulo\OneDrive\td junto outlook hotmail\MAPLE BEAR\provas\00_COORDENACAO\TRANSCRICAO_Materiais_Escola.docx'
doc.save(out)
print('OK: ' + out)
